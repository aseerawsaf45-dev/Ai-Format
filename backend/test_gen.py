import docx
from docx.oxml import parse_xml
from app.renderer.diagrams import fetch_svg_ast, parse_svg_to_ast
from app.renderer.generator import DrawingMLGenerator

doc = docx.Document()
p = doc.add_paragraph()
r = p.add_run()

mermaid_code = 'graph TD\n  A[Alice] --> B[Bob]'
svg = fetch_svg_ast(mermaid_code)
ast = parse_svg_to_ast(svg)
gen = DrawingMLGenerator(ast)
wpg_xml = gen.build_xml()

drawing_xml = f'''
<w:drawing xmlns:w="http://schemas.openxmlformats.org/wordprocessingml/2006/main" 
           xmlns:wp="http://schemas.openxmlformats.org/drawingml/2006/wordprocessingDrawing" 
           xmlns:a="http://schemas.openxmlformats.org/drawingml/2006/main" 
           xmlns:wpg="http://schemas.microsoft.com/office/word/2010/wordprocessingGroup" 
           xmlns:wps="http://schemas.microsoft.com/office/word/2010/wordprocessingShape">
    <wp:inline distT="0" distB="0" distL="0" distR="0">
        <wp:extent cx="{int(ast.width*9525)}" cy="{int(ast.height*9525)}"/>
        <wp:effectExtent l="0" t="0" r="0" b="0"/>
        <wp:docPr id="1" name="Diagram"/>
        <wp:cNvGraphicFramePr/>
        <a:graphic>
            <a:graphicData uri="http://schemas.microsoft.com/office/word/2010/wordprocessingGroup">
                {wpg_xml}
            </a:graphicData>
        </a:graphic>
    </wp:inline>
</w:drawing>
'''

r._r.append(parse_xml(drawing_xml))

mc_ignorable = doc.element.get('{http://schemas.openxmlformats.org/markup-compatibility/2006}Ignorable')
if mc_ignorable:
    if 'wpg' not in mc_ignorable: mc_ignorable += ' wpg'
    if 'wps' not in mc_ignorable: mc_ignorable += ' wps'
    doc.element.set('{http://schemas.openxmlformats.org/markup-compatibility/2006}Ignorable', mc_ignorable)

doc.save('test_generator.docx')
print('Generated test_generator.docx')
