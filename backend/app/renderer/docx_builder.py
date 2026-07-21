"""DocTree → .docx renderer — full themed implementation.

Handles every block and inline type the parser can produce:

Block types
    heading (h1–h6), paragraph, code_block, block_equation,
    blockquote, list (ordered/unordered/task, nested), table,
    image (URL + base64), thematic_break.

Inline types
    run (bold / italic / strikethrough / code), link (hyperlink),
    inline_equation (OMML fragment).

All visual choices (fonts, colours, spacing) are driven by :class:`Theme`.
"""

from __future__ import annotations

import base64
import io
import urllib.request
from io import BytesIO
from typing import cast

from docx import Document as DocxDocument
from docx.enum.text import WD_ALIGN_PARAGRAPH
from docx.oxml import OxmlElement
from docx.oxml.ns import qn
from docx.shared import Inches, Pt, RGBColor
from lxml import etree

from app.models import (
    BlockEquation,
    BlockQuote,
    CodeBlock,
    Document,
    Heading,
    Image,
    InlineEquation,
    ListBlock,
    ListItem,
    Paragraph,
    Run,
    Table,
    ThematicBreak,
)
from app.renderer.drawingml import svg_to_drawingml
from app.themes import Theme, get_theme


# ---------------------------------------------------------------------------
# Helpers — colour
# ---------------------------------------------------------------------------

def _rgb(hex6: str) -> RGBColor:
    """Parse a 6-digit hex colour (no ``#``) into an :class:`RGBColor`."""
    h = hex6.lstrip("#")
    return RGBColor(int(h[0:2], 16), int(h[2:4], 16), int(h[4:6], 16))


# ---------------------------------------------------------------------------
# Helpers — paragraph / run formatting
# ---------------------------------------------------------------------------

_ALIGN_MAP = {
    "left": WD_ALIGN_PARAGRAPH.LEFT,
    "center": WD_ALIGN_PARAGRAPH.CENTER,
    "right": WD_ALIGN_PARAGRAPH.RIGHT,
    "justify": WD_ALIGN_PARAGRAPH.JUSTIFY,
}


def _set_para_spacing(para, before_pt: float = 0, after_pt: float = 6) -> None:
    pPr = para._p.get_or_add_pPr()
    spacing = OxmlElement("w:spacing")
    spacing.set(qn("w:before"), str(int(before_pt * 20)))
    spacing.set(qn("w:after"), str(int(after_pt * 20)))
    pPr.append(spacing)


def _set_shading(cell_or_para_el, fill_hex: str) -> None:
    """Apply a solid background fill to a table cell or paragraph element."""
    if fill_hex == "transparent":
        return
    shd = OxmlElement("w:shd")
    shd.set(qn("w:val"), "clear")
    shd.set(qn("w:color"), "auto")
    shd.set(qn("w:fill"), fill_hex.lstrip("#").upper())
    cell_or_para_el.append(shd)


def _apply_run_style(run, theme: Theme, *, bold=False, italic=False,
                     strike=False, code=False) -> None:
    run.bold = bold
    run.italic = italic
    run.font.strike = strike
    if code:
        run.font.name = "Courier New"
        run.font.size = Pt(9.5)
        run.font.color.rgb = _rgb(theme.code_fg)
    else:
        run.font.name = theme.body_font
        run.font.size = Pt(theme.body_size_pt)
        run.font.color.rgb = _rgb(theme.body_color)


# ---------------------------------------------------------------------------
# Helpers — OMML injection
# ---------------------------------------------------------------------------

def _inject_omml(para, omml_fragment: str) -> None:
    """Inject a raw OMML fragment (``<m:oMath>…`` or ``<m:oMathPara>…``)
    directly into a paragraph's XML.  The paragraph element already exists;
    we just append the parsed OMML subtree to ``<w:p>``."""
    try:
        root = etree.fromstring(omml_fragment.encode())
        para._p.append(root)
    except Exception:
        # Fallback: render as plain text.
        r = para.add_run(f"[equation error]")
        r.italic = True


def _omml_para(doc: DocxDocument, omml_fragment: str) -> None:
    """Add a centred display-math paragraph by injecting ``<m:oMathPara>``."""
    para = doc.add_paragraph()
    para.alignment = WD_ALIGN_PARAGRAPH.CENTER
    _inject_omml(para, omml_fragment)
    _set_para_spacing(para, before_pt=6, after_pt=6)


# ---------------------------------------------------------------------------
# Inline rendering
# ---------------------------------------------------------------------------

def _render_inline_into_para(para, inlines: list, theme: Theme) -> None:
    """Add inline nodes to an existing paragraph."""
    for node in inlines:
        ntype = getattr(node, "type", None)
        if ntype == "run":
            r = node  # type: Run
            run = para.add_run(r.text)
            _apply_run_style(
                run, theme,
                bold=r.bold, italic=r.italic, strike=r.strike, code=r.code,
            )
        elif ntype == "link":
            # Hyperlinks need relationship + complex XML; use simple underlined run.
            run = para.add_run(node.text)
            run.font.name = theme.body_font
            run.font.size = Pt(theme.body_size_pt)
            run.font.color.rgb = _rgb(theme.link_color)
            run.font.underline = True
        elif ntype == "inline_equation":
            eq = node  # type: InlineEquation
            if eq.omml:
                _inject_omml(para, eq.omml)
            else:
                run = para.add_run(f"${eq.latex}$")
                run.italic = True
        # image nodes inside inline are rare; skip gracefully


# ---------------------------------------------------------------------------
# Block renderers
# ---------------------------------------------------------------------------

def _render_heading(doc: DocxDocument, block: Heading, theme: Theme) -> None:
    level = block.level
    para = doc.add_paragraph()
    _set_para_spacing(para, before_pt=12, after_pt=4)
    run = para.add_run(block.text)
    run.bold = True
    run.font.name = theme.heading_font
    run.font.size = Pt(theme.heading_size(level))
    run.font.color.rgb = _rgb(theme.heading_color)


def _render_paragraph(doc: DocxDocument, block: Paragraph, theme: Theme) -> None:
    para = doc.add_paragraph()
    align = _ALIGN_MAP.get(getattr(block, "align", "left"), WD_ALIGN_PARAGRAPH.LEFT)
    para.alignment = align
    _set_para_spacing(para, before_pt=0, after_pt=6)
    _render_inline_into_para(para, block.inline, theme)


def _render_code_block(doc: DocxDocument, block: CodeBlock, theme: Theme) -> None:
    """Render a fenced code block: dark-bg paragraph with monospace text."""
    if block.lang and block.lang.lower() == "mermaid":
        try:
            b64_code = base64.urlsafe_b64encode(block.code.encode('utf-8')).decode('utf-8')
            
            # Attempt to fetch and parse SVG into native DrawingML
            try:
                req_svg = urllib.request.Request(
                    f"https://mermaid.ink/svg/{b64_code}",
                    headers={"User-Agent": "Mozilla/5.0"}
                )
                with urllib.request.urlopen(req_svg, timeout=10) as resp_svg:
                    svg_str = resp_svg.read().decode('utf-8')
                
                drawing_xml = svg_to_drawingml(svg_str)
                if drawing_xml is not None:
                    # Successfully mapped to DrawingML
                    p = doc.add_paragraph()
                    p._p.append(drawing_xml)
                    
                    spacer = doc.add_paragraph()
                    _set_para_spacing(spacer, before_pt=0, after_pt=8)
                    return
            except Exception as svg_e:
                print(f"DrawingML generation failed: {svg_e}, falling back to PNG")

            # Fallback to PNG
            req = urllib.request.Request(
                f"https://mermaid.ink/img/{b64_code}",
                headers={"User-Agent": "Mozilla/5.0"}
            )
            with urllib.request.urlopen(req, timeout=10) as resp:
                stream = BytesIO(resp.read())
            doc.add_picture(stream, width=Inches(5))
            
            # Trailing spacer
            spacer = doc.add_paragraph()
            _set_para_spacing(spacer, before_pt=0, after_pt=8)
            return
        except Exception as e:
            # Fallback to standard text rendering if the API fails
            pass


    # Language label (if present).
    if block.lang:
        label_para = doc.add_paragraph()
        _set_para_spacing(label_para, before_pt=8, after_pt=0)
        lr = label_para.add_run(block.lang)
        lr.font.name = "Courier New"
        lr.font.size = Pt(8)
        lr.font.color.rgb = _rgb(theme.code_lang_color)
        pPr = label_para._p.get_or_add_pPr()
        _set_shading(pPr, theme.code_bg)

    lines = block.code.splitlines()
    for lineno, line in enumerate(lines, 1):
        p = doc.add_paragraph()
        _set_para_spacing(p, before_pt=0, after_pt=0)
        pPr = p._p.get_or_add_pPr()
        _set_shading(pPr, theme.code_bg)
        text = f"{lineno:>3}  {line}" if block.show_line_numbers else line
        run = p.add_run(text)
        run.font.name = "Courier New"
        run.font.size = Pt(9.5)
        run.font.color.rgb = _rgb(theme.code_fg)

    # Trailing spacer
    spacer = doc.add_paragraph()
    _set_para_spacing(spacer, before_pt=0, after_pt=8)


def _render_blockquote(doc: DocxDocument, block: BlockQuote, theme: Theme) -> None:
    """Render blockquote as indented paragraphs with left border shading."""
    for child in block.children:
        ctype = getattr(child, "type", None)
        if ctype == "paragraph":
            p = doc.add_paragraph()
            p.paragraph_format.left_indent = Inches(0.4)
            _set_para_spacing(p, before_pt=2, after_pt=2)
            pPr = p._p.get_or_add_pPr()
            _set_shading(pPr, theme.quote_bg)
            # Add a left border via pBdr
            pBdr = OxmlElement("w:pBdr")
            left = OxmlElement("w:left")
            left.set(qn("w:val"), "single")
            left.set(qn("w:sz"), "24")
            left.set(qn("w:space"), "4")
            left.set(qn("w:color"), theme.quote_border.lstrip("#").upper())
            pBdr.append(left)
            pPr.append(pBdr)
            _render_inline_into_para(p, child.inline, theme)
        else:
            _render_block(doc, child, theme)


def _render_list(doc: DocxDocument, block: ListBlock, theme: Theme,
                 depth: int = 0) -> None:
    """Render a list block recursively."""
    indent = Inches(0.25 + depth * 0.25)
    for idx, item in enumerate(block.items, 1):
        p = doc.add_paragraph()
        p.paragraph_format.left_indent = indent
        _set_para_spacing(p, before_pt=1, after_pt=1)

        # Bullet / number prefix.
        if block.ordered:
            prefix = f"{idx}. "
        elif item.checked is not None:
            prefix = "☑ " if item.checked else "☐ "
        else:
            bullets = ["•", "◦", "▪"]
            prefix = bullets[min(depth, 2)] + " "

        run = p.add_run(prefix)
        run.font.name = theme.body_font
        run.font.size = Pt(theme.body_size_pt)
        run.font.color.rgb = _rgb(theme.body_color)

        _render_inline_into_para(p, item.inline, theme)

        # Nested children.
        for child in item.children:
            ctype = getattr(child, "type", None)
            if ctype == "list":
                _render_list(doc, child, theme, depth=depth + 1)
            else:
                _render_block(doc, child, theme)


def _render_table(doc: DocxDocument, block: Table, theme: Theme) -> None:
    if not block.rows:
        return
    ncols = max(len(row.cells) for row in block.rows)
    tbl = doc.add_table(rows=0, cols=ncols)
    tbl.style = "Table Grid"

    for rowi, row in enumerate(block.rows):
        tr = tbl.add_row()
        is_header_row = any(c.is_header for c in row.cells)
        for ci, cell_node in enumerate(row.cells):
            cell = tr.cells[ci]
            # Background.
            tc = cell._tc
            tcPr = tc.get_or_add_tcPr()
            if is_header_row:
                _set_shading(tcPr, theme.table_header_bg)
            elif rowi % 2 == 0:
                _set_shading(tcPr, theme.table_band_bg)

            # Cell paragraph.
            para = cell.paragraphs[0]
            align_str = getattr(cell_node, "align", "left")
            para.alignment = _ALIGN_MAP.get(align_str, WD_ALIGN_PARAGRAPH.LEFT)

            for node in cell_node.inline:
                ntype = getattr(node, "type", None)
                if ntype == "run":
                    r = cell.paragraphs[0].add_run(node.text)
                    r.bold = node.bold or is_header_row
                    r.italic = node.italic
                    r.font.name = theme.body_font
                    r.font.size = Pt(theme.body_size_pt)
                    r.font.color.rgb = (
                        _rgb(theme.table_header_fg)
                        if is_header_row
                        else _rgb(theme.body_color)
                    )
                elif ntype == "link":
                    r = cell.paragraphs[0].add_run(node.text)
                    r.font.name = theme.body_font
                    r.font.size = Pt(theme.body_size_pt)
                    r.font.color.rgb = _rgb(theme.link_color)
                    r.font.underline = True


def _render_image(doc: DocxDocument, block: Image, theme: Theme) -> None:
    """Embed an image node (URL or base64) into the document."""
    try:
        if block.base64:
            data = base64.b64decode(block.base64)
            stream = BytesIO(data)
            doc.add_picture(stream, width=Inches(4))
        elif block.url:
            if block.url.startswith(("http://", "https://")):
                req = urllib.request.Request(
                    block.url,
                    headers={"User-Agent": "Mozilla/5.0"}
                )
                with urllib.request.urlopen(req, timeout=5) as resp:
                    stream = BytesIO(resp.read())
                doc.add_picture(stream, width=Inches(4))
            # Skip local/relative URLs (can't resolve at render time)
    except Exception:
        # Fallback: caption paragraph.
        p = doc.add_paragraph()
        p.alignment = WD_ALIGN_PARAGRAPH.CENTER
        r = p.add_run(f"[Image: {block.alt or block.url or 'embedded'}]")
        r.italic = True
        r.font.color.rgb = _rgb("888888")

    if block.caption:
        cap = doc.add_paragraph(block.caption)
        cap.alignment = WD_ALIGN_PARAGRAPH.CENTER
        cap.runs[0].italic = True
        cap.runs[0].font.size = Pt(9)
        _set_para_spacing(cap, before_pt=2, after_pt=8)


def _render_thematic_break(doc: DocxDocument, theme: Theme) -> None:
    """Horizontal rule — paragraph with a bottom border."""
    p = doc.add_paragraph()
    _set_para_spacing(p, before_pt=6, after_pt=6)
    pPr = p._p.get_or_add_pPr()
    pBdr = OxmlElement("w:pBdr")
    bottom = OxmlElement("w:bottom")
    bottom.set(qn("w:val"), "single")
    bottom.set(qn("w:sz"), "6")
    bottom.set(qn("w:space"), "1")
    bottom.set(qn("w:color"), theme.accent.lstrip("#").upper())
    pBdr.append(bottom)
    pPr.append(pBdr)


def _render_block_equation(doc: DocxDocument, block: BlockEquation, theme: Theme) -> None:
    if block.omml:
        _omml_para(doc, block.omml)
    else:
        # Graceful degradation: centred italic LaTeX source.
        p = doc.add_paragraph(f"$${block.latex}$$")
        p.alignment = WD_ALIGN_PARAGRAPH.CENTER
        p.runs[0].italic = True
        _set_para_spacing(p, before_pt=6, after_pt=6)


# ---------------------------------------------------------------------------
# Block dispatcher
# ---------------------------------------------------------------------------

def _render_block(doc: DocxDocument, block, theme: Theme) -> None:
    btype = getattr(block, "type", None)
    if btype == "heading":
        _render_heading(doc, block, theme)
    elif btype == "paragraph":
        _render_paragraph(doc, block, theme)
    elif btype == "code_block":
        _render_code_block(doc, block, theme)
    elif btype == "block_equation":
        _render_block_equation(doc, block, theme)
    elif btype == "blockquote":
        _render_blockquote(doc, block, theme)
    elif btype == "list":
        _render_list(doc, block, theme)
    elif btype == "table":
        _render_table(doc, block, theme)
    elif btype == "image":
        _render_image(doc, block, theme)
    elif btype == "thematic_break":
        _render_thematic_break(doc, theme)
    # Unknown types are silently skipped.


# ---------------------------------------------------------------------------
# Public entrypoint
# ---------------------------------------------------------------------------

def render_document(document: Document, theme_name: str = "modern") -> bytes:
    """Render a DocTree to ``.docx`` bytes using the named theme."""
    theme = get_theme(theme_name)
    doc = DocxDocument()

    # Page margins from theme.
    for section in doc.sections:
        m = Inches(theme.margin_inches)
        section.top_margin = m
        section.bottom_margin = m
        section.left_margin = m
        section.right_margin = m

    for block in document.children:
        _render_block(doc, block, theme)

    buf = BytesIO()
    doc.save(buf)
    return buf.getvalue()
