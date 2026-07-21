import urllib.request
import docx
from docx.opc.part import Part
from docx.shared import Inches
from io import BytesIO

doc = docx.Document()
p = doc.add_paragraph()
r = p.add_run()

b64 = 'c2VxdWVuY2VEaWFncmFtCkFsaWNlLT4+Qm9iOiBIZWxsbw=='

# 1. Add PNG picture
req_png = urllib.request.Request('https://mermaid.ink/img/' + b64, headers={'User-Agent': 'Mozilla/5.0'})
with urllib.request.urlopen(req_png) as resp:
    png_bytes = resp.read()
r.add_picture(BytesIO(png_bytes), width=Inches(5))

# 2. Get SVG bytes
req_svg = urllib.request.Request('https://mermaid.ink/svg/' + b64, headers={'User-Agent': 'Mozilla/5.0'})
with urllib.request.urlopen(req_svg) as resp:
    svg_bytes = resp.read()

# 3. Create SVG Part and relate it
partname = doc.part.package.next_partname('/word/media/image%d.svg')
svg_part = Part(partname, 'image/svg+xml', svg_bytes, doc.part.package)
rId_svg = doc.part.relate_to(svg_part, 'http://schemas.openxmlformats.org/officeDocument/2006/relationships/image')

# 4. Inject SVG extension into blip
blip = r._r.find('.//{http://schemas.openxmlformats.org/drawingml/2006/main}blip')
from docx.oxml import parse_xml
ext_xml = f'''
<a:extLst xmlns:a="http://schemas.openxmlformats.org/drawingml/2006/main">
    <a:ext uri="{{96DAC541-7B7A-43D3-8B79-37D633B846F1}}">
        <asvg:svgBlip xmlns:asvg="http://schemas.microsoft.com/office/drawing/2016/SVG/main" r:embed="{rId_svg}"/>
    </a:ext>
</a:extLst>
'''
extLst = parse_xml(ext_xml)
blip.append(extLst)

doc.save('test_svg_injection.docx')
print('Injected successfully')
