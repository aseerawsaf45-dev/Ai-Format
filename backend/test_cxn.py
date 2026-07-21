import docx
from docx.oxml import parse_xml
doc = docx.Document()
p = doc.add_paragraph()
r = p.add_run()

drawing_xml = '''
<w:drawing xmlns:w="http://schemas.openxmlformats.org/wordprocessingml/2006/main" 
           xmlns:wp="http://schemas.openxmlformats.org/drawingml/2006/wordprocessingDrawing" 
           xmlns:a="http://schemas.openxmlformats.org/drawingml/2006/main" 
           xmlns:wpg="http://schemas.microsoft.com/office/word/2010/wordprocessingGroup" 
           xmlns:wps="http://schemas.microsoft.com/office/word/2010/wordprocessingShape">
    <wp:inline distT="0" distB="0" distL="0" distR="0">
        <wp:extent cx="3000000" cy="2000000"/>
        <wp:effectExtent l="0" t="0" r="0" b="0"/>
        <wp:docPr id="1" name="Connected Diagram"/>
        <wp:cNvGraphicFramePr/>
        <a:graphic>
            <a:graphicData uri="http://schemas.microsoft.com/office/word/2010/wordprocessingGroup">
                <wpg:wgp>
                    <wpg:cNvPr id="2" name="Group"/>
                    <wpg:cNvGrpSpPr/>
                    <wpg:grpSpPr>
                        <a:xfrm>
                            <a:off x="0" y="0"/>
                            <a:ext cx="3000000" cy="2000000"/>
                            <a:chOff x="0" y="0"/>
                            <a:chExt cx="3000000" cy="2000000"/>
                        </a:xfrm>
                    </wpg:grpSpPr>
                    <!-- Shape 1 -->
                    <wps:wsp>
                        <wps:cNvPr id="3" name="Node1"/>
                        <wps:cNvSpPr/>
                        <wps:spPr>
                            <a:xfrm>
                                <a:off x="100000" y="100000"/>
                                <a:ext cx="500000" cy="500000"/>
                            </a:xfrm>
                            <a:prstGeom prst="rect"><a:avLst/></a:prstGeom>
                            <a:solidFill><a:srgbClr val="FF0000"/></a:solidFill>
                        </wps:spPr>
                        <wps:bodyPr/>
                    </wps:wsp>
                    <!-- Shape 2 -->
                    <wps:wsp>
                        <wps:cNvPr id="4" name="Node2"/>
                        <wps:cNvSpPr/>
                        <wps:spPr>
                            <a:xfrm>
                                <a:off x="2000000" y="100000"/>
                                <a:ext cx="500000" cy="500000"/>
                            </a:xfrm>
                            <a:prstGeom prst="rect"><a:avLst/></a:prstGeom>
                            <a:solidFill><a:srgbClr val="00FF00"/></a:solidFill>
                        </wps:spPr>
                        <wps:bodyPr/>
                    </wps:wsp>
                    <!-- Connector -->
                    <wps:wsp>
                        <wps:cNvPr id="5" name="Conn1"/>
                        <wps:cNvCnPr>
                            <a:cxnSpPr>
                                <a:stCxn id="3" idx="3"/>
                                <a:endCxn id="4" idx=\"1\"/>
                            </a:cxnSpPr>
                        </wps:cNvCnPr>
                        <wps:spPr>
                            <a:xfrm>
                                <a:off x="600000" y="350000"/>
                                <a:ext cx="1400000" cy="0"/>
                            </a:xfrm>
                            <a:prstGeom prst="line"><a:avLst/></a:prstGeom>
                            <a:ln w="12700">
                                <a:solidFill><a:srgbClr val="000000"/></a:solidFill>
                            </a:ln>
                        </wps:spPr>
                        <wps:bodyPr/>
                    </wps:wsp>
                </wpg:wgp>
            </a:graphicData>
        </a:graphic>
    </wp:inline>
</w:drawing>
'''
r._r.append(parse_xml(drawing_xml))

mc_ignorable = doc.element.get('{http://schemas.openxmlformats.org/markup-compatibility/2006}Ignorable')
if mc_ignorable:
    if 'wpg' not in mc_ignorable: mc_ignorable += ' wpg'
    if 'wps' not in mc_ignorable: mc_ignorable += ' wps'
    doc.element.set('{http://schemas.openxmlformats.org/markup-compatibility/2006}Ignorable', mc_ignorable)

doc.save('test_cxn.docx')
print('Created test_cxn.docx')
